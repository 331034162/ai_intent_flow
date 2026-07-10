"""
文件加载入库工具类
支持多种文件格式的加载、清洗和入库
"""
import os
import re
import uuid
import traceback
from typing import Optional, List
from dataclasses import dataclass
from datetime import datetime

from sqlalchemy import text
from .async_mysql_connection import get_async_pool_instance, AsyncMySQLConnection
from ..core.logger import app_logger


# 支持的文件类型映射
SUPPORTED_FILE_TYPES = {
    '.txt': 'txt',
    '.md': 'markdown',
    '.markdown': 'markdown',
    '.xls': 'excel',
    '.xlsx': 'excel',
    '.pdf': 'pdf',
    '.doc': 'word',
    '.docm': 'word',
    '.docx': 'word',
    '.csv': 'csv'
}


@dataclass
class UploadFile:
    """上传文件数据类"""
    file_id: str
    user_id: str
    file_name: str
    file_type: str
    file_size: int
    file_path: str
    file_content: str
    status: int = 1


class FileLoad2DB:
    """文件加载入库工具类"""

    @staticmethod
    def _get_file_type(file_path: str) -> Optional[str]:
        """
        获取文件类型

        Args:
            file_path: 文件路径

        Returns:
            文件类型，如果不支持则返回 None
        """
        ext = os.path.splitext(file_path)[1].lower()
        return SUPPORTED_FILE_TYPES.get(ext)

    @staticmethod
    def _clean_content(content: str) -> str:
        """
        清洗文件内容

        Args:
            content: 原始内容

        Returns:
            清洗后的内容
        """
        if not content:
            return ""

        # 1. 去除 BOM 字符
        content = content.replace('\ufeff', '')

        # 2. 统一换行符
        content = content.replace('\r\n', '\n').replace('\r', '\n')

        # 3. 去除多余空行（超过2个连续空行压缩为2个）
        content = re.sub(r'\n{3,}', '\n\n', content)

        # 4. 去除行首行尾空白
        lines = [line.strip() for line in content.split('\n')]
        content = '\n'.join(lines)

        # 5. 去除全角空格和特殊空白字符
        content = re.sub(r'[\u3000\x00-\x08\x0b\x0c\x0e-\x1f]', ' ', content)

        # 6. 压缩多个连续空格为单个
        content = re.sub(r' {2,}', ' ', content)

        return content.strip()

    @staticmethod
    def _load_txt_file(file_path: str) -> str:
        """加载 txt/markdown 文件"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            # 尝试其他编码
            with open(file_path, 'r', encoding='gbk') as f:
                return f.read()

    @staticmethod
    def _load_csv_file(file_path: str) -> str:
        """加载 CSV 文件"""
        import csv

        content_lines = []
        try:
            with open(file_path, 'r', encoding='utf-8', newline='') as f:
                reader = csv.reader(f)
                for row in reader:
                    content_lines.append(' | '.join(row))
        except UnicodeDecodeError:
            with open(file_path, 'r', encoding='gbk', newline='') as f:
                reader = csv.reader(f)
                for row in reader:
                    content_lines.append(' | '.join(row))

        return '\n'.join(content_lines)

    @staticmethod
    def _load_pdf_file(file_path: str) -> str:
        """加载 PDF 文件"""
        try:
            from langchain_community.document_loaders import PyPDFLoader
            loader = PyPDFLoader(file_path)
            pages = loader.load()
            return '\n\n'.join([page.page_content for page in pages])
        except ImportError:
            app_logger.warning("PyPDFLoader 不可用，尝试使用 pdfplumber")
            import pdfplumber
            content = []
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        content.append(text)
            return '\n\n'.join(content)

    @staticmethod
    def _load_word_file(file_path: str) -> str:
        """加载 Word 文件"""
        try:
            from langchain_community.document_loaders import Docx2txtLoader
            loader = Docx2txtLoader(file_path)
            docs = loader.load()
            return '\n\n'.join([doc.page_content for doc in docs])
        except ImportError:
            app_logger.warning("Docx2txtLoader 不可用，尝试使用 python-docx")
            from docx import Document
            doc = Document(file_path)
            content = []
            for para in doc.paragraphs:
                if para.text.strip():
                    content.append(para.text)
            # 提取表格内容
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join([cell.text.strip() for cell in row.cells])
                    if row_text.strip():
                        content.append(row_text)
            return '\n\n'.join(content)

    @staticmethod
    def _load_excel_file(file_path: str) -> str:
        """加载 Excel 文件"""
        try:
            from langchain_community.document_loaders import UnstructuredExcelLoader
            loader = UnstructuredExcelLoader(file_path)
            docs = loader.load()
            return '\n\n'.join([doc.page_content for doc in docs])
        except ImportError:
            app_logger.warning("UnstructuredExcelLoader 不可用，尝试使用 pandas")
            import pandas as pd
            # 读取所有 sheet
            excel_file = pd.ExcelFile(file_path)
            content = []
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(excel_file, sheet_name=sheet_name)
                content.append(f"=== Sheet: {sheet_name} ===")
                # 转换为表格格式
                content.append(df.to_string(index=False))
            return '\n\n'.join(content)

    @staticmethod
    def load_file_content(file_path: str) -> Optional[str]:
        """
        根据文件类型加载文件内容

        Args:
            file_path: 文件路径

        Returns:
            文件内容字符串，加载失败返回 None
        """
        file_type = FileLoad2DB._get_file_type(file_path)
        if not file_type:
            app_logger.error(f"不支持的文件类型: {file_path}")
            return None

        try:
            if file_type in ('txt', 'markdown'):
                return FileLoad2DB._load_txt_file(file_path)
            elif file_type == 'csv':
                return FileLoad2DB._load_csv_file(file_path)
            elif file_type == 'pdf':
                return FileLoad2DB._load_pdf_file(file_path)
            elif file_type == 'word':
                return FileLoad2DB._load_word_file(file_path)
            elif file_type == 'excel':
                return FileLoad2DB._load_excel_file(file_path)
            else:
                app_logger.error(f"未知文件类型: {file_type}")
                return None
        except Exception as e:
            app_logger.error(f"加载文件失败: {file_path}, 错误: {e}\n{traceback.format_exc()}")
            return None

    @staticmethod
    async def load_and_save(
        user_id: str,
        file_path: str
    ) -> Optional[str]:
        """
        加载文件内容并保存到数据库

        Args:
            user_id: 用户ID
            file_path: 文件路径

        Returns:
            成功返回 file_id，失败返回 None
        """
        # 检查文件是否存在
        if not os.path.exists(file_path):
            app_logger.error(f"文件不存在: {file_path}")
            return None

        # 获取文件类型
        file_type = FileLoad2DB._get_file_type(file_path)
        if not file_type:
            app_logger.error(f"不支持的文件类型: {file_path}")
            return None

        # 加载文件内容
        raw_content = FileLoad2DB.load_file_content(file_path)
        if raw_content is None:
            return None

        # 清洗内容
        cleaned_content = FileLoad2DB._clean_content(raw_content)

        # 获取文件信息
        file_name = os.path.basename(file_path)
        file_size = os.path.getsize(file_path)
        file_id = str(uuid.uuid4())

        try:
            # 保存到数据库
            db_conn = await get_async_pool_instance()
            session = await db_conn.get_session()

            insert_sql = text("""
                INSERT INTO zb_conversation_upload_files (
                    file_id, user_id, file_name, file_type,
                    file_size, file_path, file_content, status, create_time
                ) VALUES (
                    :file_id, :user_id, :file_name, :file_type,
                    :file_size, :file_path, :file_content, :status, :create_time
                )
            """)

            params = {
                "file_id": file_id,
                "user_id": user_id,
                "file_name": file_name,
                "file_type": file_type,
                "file_size": file_size,
                "file_path": file_path,
                "file_content": cleaned_content,
                "status": 1,
                "create_time": datetime.now()
            }

            async with session:
                async with session.begin():
                    await session.execute(insert_sql, params)

            app_logger.info(
                f"文件入库成功 | file_id={file_id} | user_id={user_id} | "
                f"file_name={file_name} | file_type={file_type} | "
                f"file_size={file_size} | content_length={len(cleaned_content)}"
            )

            return file_id

        except Exception as e:
            app_logger.error(f"文件入库失败: {e}\n{traceback.format_exc()}")
            return None

    @staticmethod
    async def get_file_by_id(file_id: str) -> Optional[UploadFile]:
        """
        根据 file_id 获取文件信息

        Args:
            file_id: 文件ID

        Returns:
            UploadFile 对象，未找到返回 None
        """
        try:
            db_conn = await get_async_pool_instance()
            session = await db_conn.get_session()

            query = text("""
                SELECT * FROM zb_conversation_upload_files
                WHERE file_id = :file_id AND status = 1
                LIMIT 1
            """)

            async with session:
                result = AsyncMySQLConnection.one(
                    await session.execute(query, {"file_id": file_id})
                )

                if result:
                    return UploadFile(
                        file_id=result['file_id'],
                        user_id=result['user_id'],
                        file_name=result['file_name'],
                        file_type=result['file_type'],
                        file_size=result['file_size'],
                        file_path=result['file_path'],
                        file_content=result['file_content'],
                        status=result['status']
                    )

            return None

        except Exception as e:
            app_logger.error(f"获取文件信息失败: {e}\n{traceback.format_exc()}")
            return None

    @staticmethod
    async def get_files_by_user(user_id: str) -> List[UploadFile]:
        """
        获取用户的所有文件

        Args:
            user_id: 用户ID

        Returns:
            UploadFile 列表
        """
        try:
            db_conn = await get_async_pool_instance()
            session = await db_conn.get_session()

            query = text("""
                SELECT * FROM zb_conversation_upload_files
                WHERE user_id = :user_id AND status = 1
                ORDER BY create_time DESC
            """)

            async with session:
                results = AsyncMySQLConnection.all(
                    await session.execute(query, {"user_id": user_id})
                )

                return [
                    UploadFile(
                        file_id=r['file_id'],
                        user_id=r['user_id'],
                        file_name=r['file_name'],
                        file_type=r['file_type'],
                        file_size=r['file_size'],
                        file_path=r['file_path'],
                        file_content=r['file_content'],
                        status=r['status']
                    )
                    for r in results
                ]

        except Exception as e:
            app_logger.error(f"获取用户文件列表失败: {e}\n{traceback.format_exc()}")
            return []

    @staticmethod
    async def get_files_by_ids(file_ids: List[str]) -> List[UploadFile]:
        """
        根据 file_id 列表批量获取文件信息

        Args:
            file_ids: 文件ID列表

        Returns:
            UploadFile 列表
        """
        if not file_ids:
            return []

        try:
            db_conn = await get_async_pool_instance()
            session = await db_conn.get_session()

            placeholders = ', '.join([f':file_id_{i}' for i in range(len(file_ids))])
            params = {f'file_id_{i}': fid for i, fid in enumerate(file_ids)}

            query = text(f"""
                SELECT * FROM zb_conversation_upload_files
                WHERE file_id IN ({placeholders}) AND status = 1
            """)

            async with session:
                results = AsyncMySQLConnection.all(
                    await session.execute(query, params)
                )

                return [
                    UploadFile(
                        file_id=r['file_id'],
                        user_id=r['user_id'],
                        file_name=r['file_name'],
                        file_type=r['file_type'],
                        file_size=r['file_size'],
                        file_path=r['file_path'],
                        file_content=r['file_content'],
                        status=r['status']
                    )
                    for r in results
                ]

        except Exception as e:
            app_logger.error(f"批量获取文件信息失败: {e}\n{traceback.format_exc()}")
            return []

    @staticmethod
    async def delete_file(file_id: str) -> bool:
        """
        软删除文件（将 status 置为 0）

        Args:
            file_id: 文件ID

        Returns:
            成功返回 True，失败返回 False
        """
        try:
            db_conn = await get_async_pool_instance()
            session = await db_conn.get_session()

            update_sql = text("""
                UPDATE zb_conversation_upload_files
                SET status = 0, update_time = NOW()
                WHERE file_id = :file_id
            """)

            rowcount = 0
            async with session:
                async with session.begin():
                    result = await session.execute(update_sql, {"file_id": file_id})
                    rowcount = result.rowcount

            if rowcount > 0:
                app_logger.info(f"文件已删除: file_id={file_id}")
                return True
            return False

        except Exception as e:
            app_logger.error(f"删除文件失败: {e}\n{traceback.format_exc()}")
            return False